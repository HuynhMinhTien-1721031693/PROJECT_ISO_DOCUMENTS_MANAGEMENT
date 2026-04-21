from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt


@dataclass(frozen=True)
class TaskItem:
    title: str
    evidence: str
    points: int


@dataclass(frozen=True)
class MemberReport:
    student_id: str
    full_name: str
    role: str
    output_name: str
    tasks: list[TaskItem]


REPORTS: list[MemberReport] = [
    MemberReport(
        student_id="1721031670",
        full_name="Nguyen Cong Anh",
        role="Database",
        output_name="BaoCao_1721031670_Database.docx",
        tasks=[
            TaskItem("Thiet ke truy van tim kiem co phan trang", "src/02_Application/IsoDoc.Application/Documents/Queries/GetDocuments/GetDocumentsQuery.cs", 10),
            TaskItem("Chuan hoa repository contract cho truy van tim kiem", "src/01_Domain/IsoDoc.Domain/Interfaces/IRepositories.cs", 10),
            TaskItem("Mo rong dich vu Elasticsearch cho bo loc va sap xep", "src/03_Infrastructure/IsoDoc.Infrastructure/Search/ElasticsearchService.cs", 10),
            TaskItem("Bo sung model bo loc tim kiem tai tang Infrastructure", "src/03_Infrastructure/IsoDoc.Infrastructure/Search/DocumentSearchFilter.cs", 10),
            TaskItem("Dong bo mapping du lieu tim kiem sang summary DTO", "src/02_Application/IsoDoc.Application/Documents/Mapping/DocumentSummaryMapping.cs", 10),
            TaskItem("Cap nhat cau hinh ket noi SQL/Elasticsearch cho API", "src/04_WebAPI/IsoDoc.WebAPI/appsettings.json", 10),
            TaskItem("Toi uu cau hinh Docker cho SQL Server va Elasticsearch", "docker-compose.yml, src/04_WebAPI/IsoDoc.WebAPI/appsettings.Docker.json", 10),
            TaskItem("Ra soat va cap nhat huong dan trien khai Docker", "DOCKER.md", 10),
            TaskItem("Dong bo cau hinh moi truong mau cho DB local", ".env.example, src/04_WebAPI/IsoDoc.WebAPI/appsettings.Development.json.example", 10),
            TaskItem("Xac thuc pipeline CI build/test sau thay doi DB stack", ".github/workflows/ci.yml, IsoDocumentManagement.CI.slnf", 10),
        ],
    ),
    MemberReport(
        student_id="1822041903",
        full_name="Tran Minh Tam",
        role="FE",
        output_name="BaoCao_1822041903_FE.docx",
        tasks=[
            TaskItem("Nang cap man hinh dang nhap va xu ly trang thai", "src/05_Frontend/IsoDoc.Blazor/Components/Pages/Login.razor", 10),
            TaskItem("Xay dung giao dien danh sach tai lieu co loc", "src/05_Frontend/IsoDoc.Blazor/Components/Pages/Documents.razor", 10),
            TaskItem("Bo sung trang tao tai lieu moi", "src/05_Frontend/IsoDoc.Blazor/Components/Pages/NewDocument.razor", 10),
            TaskItem("Cap nhat dashboard theo du lieu tong hop", "src/05_Frontend/IsoDoc.Blazor/Components/Pages/Dashboard.razor", 10),
            TaskItem("Chinh sua component filter panel de tim kiem nang cao", "src/05_Frontend/IsoDoc.Blazor/Components/Documents/FilterPanel.razor", 10),
            TaskItem("Chuan hoa model du lieu hien thi tai lieu", "src/05_Frontend/IsoDoc.Blazor/Models/DocumentModels.cs", 10),
            TaskItem("Mo rong service goi API tai lieu", "src/05_Frontend/IsoDoc.Blazor/Services/Api/DocumentApiService.cs", 10),
            TaskItem("Tich hop thong bao realtime SignalR cho giao dien", "src/05_Frontend/IsoDoc.Blazor/Services/Api/NotificationSignalRService.cs", 10),
            TaskItem("Dong bo tham so runtime Frontend trong docker-compose", "docker-compose.yml", 10),
            TaskItem("Bo sung screenshot minh hoa luong giao dien", "docs/screenshots/login-demo.svg, docs/screenshots/documents-demo.svg, docs/screenshots/workflow-demo.svg", 10),
        ],
    ),
    MemberReport(
        student_id="1822040198",
        full_name="Nguyen Do Khoi Nguyen",
        role="Tester",
        output_name="BaoCao_1822040198_Tester.docx",
        tasks=[
            TaskItem("Viet bo test workflow domain draft -> review -> publish", "src/tests/IsoDoc.Domain.Tests/DocumentWorkflowTests.cs", 10),
            TaskItem("Kiem thu truong hop submit khong co file", "src/tests/IsoDoc.Domain.Tests/DocumentWorkflowTests.cs", 10),
            TaskItem("Kiem thu loi workflow khi transition sai trang thai", "src/tests/IsoDoc.Domain.Tests/DocumentWorkflowTests.cs", 10),
            TaskItem("Kiem thu duong dan phe duyet day du va tang version", "src/tests/IsoDoc.Domain.Tests/DocumentWorkflowTests.cs", 10),
            TaskItem("Kiem thu luong reject va return to draft", "src/tests/IsoDoc.Domain.Tests/DocumentWorkflowTests.cs", 10),
            TaskItem("Khoi tao test project domain voi csproj rieng", "src/tests/IsoDoc.Domain.Tests/IsoDoc.Domain.Tests.csproj", 10),
            TaskItem("Cap nhat smoke test Playwright cho e2e", "e2e/playwright.config.ts", 10),
            TaskItem("Dong bo CI de chay restore/build/test tu dong", ".github/workflows/ci.yml", 10),
            TaskItem("Tao solution filter cho luong CI khong phu thuoc MAUI", "IsoDocumentManagement.CI.slnf", 10),
            TaskItem("Xac nhan log test run va ket qua lan chay gan nhat", "test-results/.last-run.json", 10),
        ],
    ),
    MemberReport(
        student_id="1822040073",
        full_name="Vu Dao Tien Phat",
        role="Tester",
        output_name="BaoCao_1822040073_Tester.docx",
        tasks=[
            TaskItem("Ra soat testcase cho API document query va bo loc", "src/02_Application/IsoDoc.Application/Documents/Queries/GetDocuments/GetDocumentsQuery.cs", 10),
            TaskItem("Danh gia tinh dung dan validation tim kiem", "src/02_Application/IsoDoc.Application/Documents/Queries/SearchDocuments/SearchDocumentsQuery.cs", 10),
            TaskItem("Kiem thu event handler dong bo chi muc tim kiem", "src/02_Application/IsoDoc.Application/Documents/EventHandlers/DocumentEventHandlers.cs", 10),
            TaskItem("Kiem thu endpoint documents va response model", "src/04_WebAPI/IsoDoc.WebAPI/Controllers/DocumentsController.cs", 10),
            TaskItem("Kiem thu endpoint workflow/auth/search tong hop", "src/04_WebAPI/IsoDoc.WebAPI/Controllers/OtherControllers.cs", 10),
            TaskItem("Kiem thu middleware va cau hinh service WebAPI", "src/04_WebAPI/IsoDoc.WebAPI/Program.cs, src/04_WebAPI/IsoDoc.WebAPI/Extensions/WebApiServiceExtensions.cs", 10),
            TaskItem("Kiem thu xac thuc local config user + lockout", "src/03_Infrastructure/IsoDoc.Infrastructure/Identity/ConfigFileUserAuthenticationService.cs", 10),
            TaskItem("Kiem thu thong bao realtime hub va notification service", "src/03_Infrastructure/IsoDoc.Infrastructure/Notifications/NotificationHub.cs, src/03_Infrastructure/IsoDoc.Infrastructure/Notifications/NotificationService.cs", 10),
            TaskItem("Cap nhat huong dan su dung va kich ban test thu cong", "HUONG_DAN_SU_DUNG.md, README.md", 10),
            TaskItem("Tong hop log loi/chay API de doi chieu ket qua test", "api-run-err.txt, api-run-out.txt", 10),
        ],
    ),
    MemberReport(
        student_id="1721031693",
        full_name="Huynh Minh Tien",
        role="Fullstack",
        output_name="BaoCao_1721031693_Fullstack.docx",
        tasks=[
            TaskItem("Thiet ke va ket noi truy van tai lieu tu Domain den Application", "src/01_Domain/IsoDoc.Domain/Interfaces/IRepositories.cs, src/02_Application/IsoDoc.Application/Common/Interfaces/IApplicationServices.cs", 10),
            TaskItem("Trien khai use case danh sach tai lieu va mapping DTO", "src/02_Application/IsoDoc.Application/Documents/Queries/GetDocuments/GetDocumentsQuery.cs, src/02_Application/IsoDoc.Application/Documents/Mapping/DocumentSummaryMapping.cs", 10),
            TaskItem("Phat trien event handler dong bo tim kiem va thong bao", "src/02_Application/IsoDoc.Application/Documents/EventHandlers/DocumentEventHandlers.cs", 10),
            TaskItem("Nang cap ha tang tim kiem Elasticsearch + bo loc", "src/03_Infrastructure/IsoDoc.Infrastructure/Search/ElasticsearchService.cs, src/03_Infrastructure/IsoDoc.Infrastructure/Search/DocumentSearchFilter.cs", 10),
            TaskItem("Trien khai identity local + bao mat dang nhap", "src/03_Infrastructure/IsoDoc.Infrastructure/Identity/ConfigFileUserAuthenticationService.cs", 10),
            TaskItem("Xay dung endpoint WebAPI cho documents, workflow, auth", "src/04_WebAPI/IsoDoc.WebAPI/Controllers/DocumentsController.cs, src/04_WebAPI/IsoDoc.WebAPI/Controllers/OtherControllers.cs", 10),
            TaskItem("To chuc startup WebAPI va DI cho cac tang", "src/04_WebAPI/IsoDoc.WebAPI/Program.cs, src/04_WebAPI/IsoDoc.WebAPI/Extensions/WebApiServiceExtensions.cs", 10),
            TaskItem("Phat trien giao dien Blazor va ket noi API/SignalR", "src/05_Frontend/IsoDoc.Blazor/Components/Pages/*.razor, src/05_Frontend/IsoDoc.Blazor/Services/Api/*.cs", 10),
            TaskItem("Dong bo van hanh Docker + tai lieu huong dan", "docker-compose.yml, DOCKER.md, HUONG_DAN_SU_DUNG.md, README.md", 10),
            TaskItem("Thiet lap CI va bo test domain/e2e cho quality gate", ".github/workflows/ci.yml, src/tests/IsoDoc.Domain.Tests/DocumentWorkflowTests.cs, e2e/playwright.config.ts", 10),
        ],
    ),
]


def create_report(base_dir: Path, report: MemberReport) -> None:
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    document.add_heading("BAO CAO PHAN CONG CONG VIEC THEO THANH VIEN", level=1)
    document.add_paragraph(f"MSSV: {report.student_id}")
    document.add_paragraph(f"Ho va ten: {report.full_name}")
    document.add_paragraph(f"Vai tro: {report.role}")
    document.add_paragraph("Nguon du lieu: tong hop tu cac tep thay doi va tai lieu hien co trong repository.")
    document.add_paragraph("Nguyen tac cham diem: tong diem can bang giua cac thanh vien (100 diem/nguoi).")

    document.add_heading("Danh sach cong viec", level=2)
    total = 0
    for index, task in enumerate(report.tasks, start=1):
        total += task.points
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(f"{task.title} ")
        paragraph.add_run(f"(Diem: {task.points})").bold = True
        document.add_paragraph(f"Minh chung: {task.evidence}")

    document.add_heading("Tong ket diem", level=2)
    document.add_paragraph(f"Tong so dau viec: {len(report.tasks)}")
    document.add_paragraph(f"Tong diem: {total}")
    document.add_paragraph("Danh gia: Dat yeu cau can bang diem so trong nhom.")

    output_path = base_dir / report.output_name
    document.save(output_path)


def main() -> None:
    docs_dir = Path(__file__).resolve().parent
    for report in REPORTS:
        create_report(docs_dir, report)
    print("Generated reports:")
    for report in REPORTS:
        print(f"- {report.output_name}")


if __name__ == "__main__":
    main()
